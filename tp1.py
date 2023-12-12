from functools import reduce
import os
from xml.dom.minidom import Document
import pandas as pd 
import datetime
import docx 
from docx import Document
import tkinter as tk
from tkinter import filedialog
import shutil


# Uploader un fichier
def upload_file(file_path):
    root, file_ext = os.path.splitext(file_path)
    if file_ext == ".txt":
        with open(file_path, "r") as f:
            data = f.read()
        return data
    else: 
        document = docx.Document(file_path)
        text = ''
        for paragraph in document.paragraphs:
            text += paragraph.text
        return text

def create_datastructure(date_table, list_files):

    # Create a dictionary to store the data for the DataFrame.
    data = {}
    for i in range(len(date_table)):
        data[date_table[i]] = list_files[i]

    return data

# Indexation automatique par date , heure : 

date_table = []
list_files = []
def index_file_by_date(file_path): 
    
    f_data = upload_file(file_path)

    # Obtenir la date et l'heure du fichier
    date_time = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
    #current_date = datetime.now()

    if date_time.date() in date_table:
        date_indx = date_table.index(date_time.date())
        list_files[date_indx].append(file_path)
    else:
        date_table.append(date_time.date())
        list_files.append([file_path])
        
    print("date_table: " , date_table)
    print("list_files: " , list_files)
    df_date = create_datastructure(date_table,list_files)
    return df_date

time_table = []
list_files_time = []
def index_file_by_time(file_path):
    
    f_data = upload_file(file_path)

    # Obtenir la date et l'heure du fichier
    date_time = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
    #current_date = datetime.now()

    if date_time.time().hour in time_table:
        time_indx = time_table.index(date_time.time().hour)
        list_files_time[time_indx].append(file_path)
    else:
        time_table.append(date_time.time().hour)
        list_files_time.append([file_path])


    #df_time = pd.DataFrame(list_files_time, index=time_table, columns=['Postings'])
    df_time = create_datastructure(time_table,list_files_time)
    return df_time
    
    

owner_table = []
owner_table_files = []

def index_file_by_owner(file_path):
    
    f_data = upload_file(file_path)

    # Obtenir le owner du ficheir 
    stat = os.stat(file_path)
    file_owner = stat.st_uid
    #current_date = datetime.now()

    if file_owner in owner_table:
        owner_indx = owner_table.index(file_owner)
        owner_table_files[owner_indx].append(file_path)
    else:
        owner_table.append(file_owner)
        owner_table_files.append([file_path])


    #df_time = pd.DataFrame(list_files_time, index=time_table, columns=['Postings'])
    df_owner = create_datastructure(owner_table,owner_table_files)
    return df_owner

type_table = []
type_table_files = []
def index_file_by_type(file_path):
    
    f_data = upload_file(file_path)

    # Obtenir le type du ficheir 
    root, file_ext = os.path.splitext(file_path)
    #current_date = datetime.now()

    if file_ext in type_table:
        ext_indx = type_table.index(file_ext)
        type_table_files[ext_indx].append(file_path)
    else:
        type_table.append(file_ext)
        type_table_files.append([file_path])


    #df_time = pd.DataFrame(list_files_time, index=time_table, columns=['Postings'])
    df_type= create_datastructure(type_table,type_table_files)
    return df_type


type_owner_table = []
type_owner_table_files = []
def index_file_by_type_owner(file_path):
    
    f_data = upload_file(file_path)

    # Obtenir le owner du ficheir 
    root, file_ext = os.path.splitext(file_path)
    stat = os.stat(file_path)
    file_owner = stat.st_uid
    #current_date = datetime.now()

    if [file_ext,file_owner] in type_owner_table:
        ext_owner_indx = type_owner_table.index([file_ext,file_owner])
        type_owner_table_files[ext_owner_indx].append(file_path)
    else:
        type_owner_table.append([file_ext,file_owner])
        type_owner_table_files.append([file_path])


    #df_time = pd.DataFrame(list_files_time, index=time_table, columns=['Postings'])
    # df_type_owner= create_datastructure(type_owner_table,type_owner_table_files)
    # return df_type_owner


#works on docs files 
type_keyword = []
type_keyword_files = []
def index_file_by_keyword(file_path):
    
    f_data = upload_file(file_path)

    # Obtenir les tags du ficheir 
    doc = Document(file_path)

    # Access document properties
    properties = doc.core_properties

    # Get keywords or tags
    tags = properties.keywords
    keywords_list = tags.split('; ')
     
    #print(tags)
    for tag in keywords_list:
        if tag in type_keyword:
            ext_indx = type_keyword.index(tag)
            type_keyword_files[ext_indx].append(file_path)
        else:
            type_keyword.append(tag)
            type_keyword_files.append([file_path])
    # print(type_keyword)
    # print(type_keyword_files)
    

# manual form : type & owner 
def index_form(type="" , owner=""  ):
    if type == "" and owner == "" :
        for i in len(type_owner_table):
            print(type_owner_table[i],": ",type_owner_table_files[i])
        return type_owner_table_files
    elif owner == "":
        type_indx = type_table.index(type)
        return type_table_files [type_indx]
    elif type == "":
        owner_indx = owner_table.index(owner)
        return owner_table_files [owner_indx]
    else:
        type_indx = type_table.index(type)
        t_type= set(type_table_files [type_indx])
        owner_indx = owner_table.index(owner)
        t_owner = set(owner_table_files [owner_indx])
        return list(t_owner & t_type)
        
import regex
import nltk

# Load the lemmatizer
lemmatizer = nltk.stem.WordNetLemmatizer()

nltk.download('stopwords')
nltk.download('wordnet')

# Load the list of stop words
stopwords = nltk.corpus.stopwords.words('english')

# Define a function to read from multiple files
def read_files(filenames):
    text = ''
    for filename in filenames:
        with open(filename, 'r') as f:
            text += f.read()
    return text.lower()

# Define a function to get the unique words and their frequency in all files
def get_unique_words_and_frequency(text):
    tokens = regex.split(r'\W+', text)
    tokens = [token for token in tokens if token not in stopwords]
    tokens = [lemmatizer.lemmatize(token) for token in tokens]
    unique_words = set(tokens)
    word_counts = {word: tokens.count(word) for word in unique_words }
    return unique_words, word_counts

# Define a function to get the files and frequency for each word
def get_files_and_frequency_for_each_word(unique_words, text, filenames):
    words_and_files = {}
    for word in unique_words:
        words_and_files[word] = {}
        for filename in filenames:
            file_text = read_files([filename])
            file_tokens = regex.split(r'\W+', file_text)
            file_stemmed_tokens = [lemmatizer.lemmatize(token) for token in file_tokens]
            file_word_counts = {word: file_stemmed_tokens.count(word) for word in unique_words}
            words_and_files[word][filename] = file_word_counts[word]
    return words_and_files

# Get the list of filenames
filenames = ["file1.txt", "file2.txt"]

# Read the text from all files
text = read_files(filenames)

# Get the unique words and their frequency in all files
unique_words, word_counts = get_unique_words_and_frequency(text)

# Get the files and frequency for each word
words_and_files = get_files_and_frequency_for_each_word(unique_words, text, filenames)

index_table = []
index_freq_table = []
index_files_table = []
def indexation ():
    
    for word, files_and_frequencies in words_and_files.items():
        index_table.append(word)
        #file_name, value = list(files_and_frequencies.items())
        file_list = []
        val_tot = 0
        for i  in range (len(list(files_and_frequencies.items()))):
            file_name , val = list(files_and_frequencies.items())[i]
            val_tot += val
            file_list.append(file_name)
        index_files_table.append(file_list)
        index_freq_table.append(val_tot)
        # print(f'{word}: {files_and_frequencies}')
    
    return index_table , index_freq_table , index_files_table
    

def choose_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window

    try:
        # Ask the user to select a file
        file_path = filedialog.askopenfilename(title="Select a file to upload")
    except FileNotFoundError:
            print("File not selected.")
    except Exception as e:
        print(f"An error occurred: {e}")
    return file_path

# création 
def arborescence(file_path):
    file_name = os.path.basename(file_path).split(".")[0]

    if not os.path.exists("arborescence"):
        os.makedirs("arborescence")

    if not os.path.exists("arborescence/SRI"):
        os.makedirs("arborescence/SRI")
    
    if not os.path.exists("arborescence/SRI/enonce"):
        os.makedirs("arborescence/SRI/enonce")

    if not os.path.exists("arborescence/SRI/CR"):
        os.makedirs("arborescence/SRI/CR")

    if not os.path.exists("arborescence/Autre"):
        os.makedirs("arborescence/Autre")
    # ajout des fichiers : 
    if('sri' in file_name.lower()):
        if('compte-rendu' in file_name.lower()):
            shutil.move(file_path, 'arborescence/SRI/CR')
        elif('enonce' in file_name.lower()):
            shutil.move(file_path, 'arborescence/SRI/enonce')
    else:
        shutil.move(file_path, 'arborescence/Autre')

            
#---------------------------------------------------------------------

# Test d'exécution de l'indexation et puis évaluation par rapport au résultat parfait
        
# file_path_1 = choose_file() # "file1.txt"
# file_path_2 = choose_file() # "file2.txt"
# file_path_3 = choose_file() # "file3.docx"

# index_file_by_type(file_path_1)
# index_file_by_type(file_path_2)
# df1 = index_file_by_type(file_path_3)

# index_file_by_owner(file_path_1)
# index_file_by_owner(file_path_2)
# df2 = index_file_by_owner(file_path_3)

# index_file_by_type_owner(file_path_1)
# index_file_by_type_owner(file_path_2)
# df3 = index_file_by_type_owner(file_path_3)

# result = index_form(".txt")
# print(result)


# recall_parfait = 1
# resultat_parfait = ['file1.txt', 'file2.txt']
# total_pertinents = len(resultat_parfait)
# truePositives = 0

# for i in range(len(result)):
#     if(resultat_parfait[i]==result[i]):
#         truePositives=truePositives+1
# print("Perfect recall is: ",recall_parfait," and obtained recall is ",truePositives/total_pertinents)

# df_date_1 = index_file_by_type_owner(file_path_1)
# df_date_2 = index_file_by_date(file_path_2)

index_t , index_freq_t , index_files_t= indexation ()
print(index_t)

# print(df_date_2)

# # Plan d'arborescence
# fileName = choose_file()
# arborescence(fileName)

# #test the tags function : 
# file_path_3 = choose_file()
# index_file_by_keyword(file_path_3)
#------------------------------------------------------------------------------------------

# search by keyword : 
def rech_mot_clé(motcle):
    if motcle in type_keyword : 
        MC_indx = type_keyword.index(motcle)
        return type_keyword_files[MC_indx]
    else : 
        print("pas de fichier correspondant à ce tag")

doc_tables = []
# search by keyword association : 
def rech_mot_clé_asso(motcles):
    mc_list = motcles.split()
    
    for mc in mc_list:
        ext_indx = type_keyword.index(mc)
        if ext_indx >=0:
            doc_tables.append(type_keyword_files[ext_indx])
       
    intersection_result = reduce(lambda x, y: set(x) & set(y), doc_tables)
    return intersection_result

#search in the content : 
def rech_contenu(sentence):
    result_table = []
    #Stemming : 
    file_tokens = regex.split(r'\W+', sentence)
    file_stemmed_tokens = [lemmatizer.lemmatize(token) for token in file_tokens]
    for word in file_stemmed_tokens :
        ext_indx = index_table.index(word)
        if ext_indx >=0 : 
            result_table.append(index_files_table[ext_indx])
    print("result_table",result_table)
    contenu_result = reduce(lambda x, y: set(x) & set(y), result_table)
    print("contenu_result : ",contenu_result)
    return contenu_result , file_stemmed_tokens
    
#evaluation metrics :
def evaluate_count(cont_res , norm_request ):
    freq_files_wd = 0
    word_sum = 0
    text_tot = read_files(cont_res)
    file_tokens = regex.split(r'\W+', text_tot)
    initial_file_stemmed_tokens = [lemmatizer.lemmatize(token) for token in file_tokens]
    word_sum += len(initial_file_stemmed_tokens)
 
    for wd in norm_request : 
        ext_indx = index_table.index(wd)
        if ext_indx >=0 : 
            freq_files_wd += index_freq_t[ext_indx]  
    
    print("word_sum",word_sum) 
    print("freq_files_wd",freq_files_wd)       
    if(word_sum != 0) :
        return  freq_files_wd / word_sum

#-------------------------------------------------------
# test the tags function : 
# file_path_3 = choose_file()
# index_file_by_keyword(file_path_3)
# # test the tags function : 
# file_path_4 = choose_file()
# index_file_by_keyword(file_path_4)
# print(type_keyword)
#print(rech_mot_clé_asso("cats dogs"))
con_res , sentence_mod = rech_contenu("cats nice")
print("mod req" , sentence_mod )
print(evaluate_count(con_res , sentence_mod ))
#-------------------------------------------------------

